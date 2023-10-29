import csv
import sqlite3
import docx
import os

# you may key in the input database file, the name of the table, and output word file here
INPUT_DATABASE_FILE = './sample-database.db'
INPUT_DATABASE_FILE_TABLE_NAME = 'book'
OUTPUT_WORD_FILE = './data.docx'


# connect to database
conn = sqlite3.connect(INPUT_DATABASE_FILE)
cursor = conn.cursor()
cursor.execute(f"select * from {INPUT_DATABASE_FILE_TABLE_NAME};")

# export database to temp csv
with open("temp_data.csv", "w", newline='') as csv_file:
    csv_writer = csv.writer(csv_file)
    csv_writer.writerow([i[0] for i in cursor.description])
    csv_writer.writerows(cursor)
conn.close()

# export temp csv to docx
doc = docx.Document()

with open('temp_data.csv', newline="") as target_file:
    csv_reader = csv.reader(target_file)

    csv_headers = next(csv_reader)
    csv_cols = len(csv_headers)

    table = doc.add_table(rows=1, cols=csv_cols)
    hdr_cells = table.rows[0].cells

    for i in range(csv_cols):
        hdr_cells[i].text = csv_headers[i]
    
    for row in csv_reader:
        row_cells = table.add_row().cells
        for i in range(csv_cols):
            row_cells[i].text = row[i]
    
    # make header bold and all caps
    for column in range(csv_cols):
        table.rows[0].cells[column].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[column].paragraphs[0].runs[0].font.all_caps = True
    
    doc.add_page_break()
    doc.save(OUTPUT_WORD_FILE)

# remove temp csv
os.remove("temp_data.csv")